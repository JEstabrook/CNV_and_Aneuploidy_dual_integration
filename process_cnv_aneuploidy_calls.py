
import os
import pandas as pd
import numpy as np
import argparse
import json
import docx
from docx.shared import Pt
import sys
import functools
import inspect
from datetime import datetime
import warnings

warnings.simplefilter(action='ignore', category=DeprecationWarning)

class PrintLogger:
    def __init__(self, log_file):
        self.log_file = log_file
        self.original_stdout = sys.stdout

    def __enter__(self):
        sys.stdout = open(self.log_file, 'a')
        return self

    def __exit__(self, exc_type, exc_value, traceback):
        sys.stdout.close()
        sys.stdout = self.original_stdout

def print_parameters_on_call(func):
    """A decorator to print the parameters passed to a function when it's called."""
    @functools.wraps(func)  # this keeps the name and docstring of the decorated function
    def wrapper(*args, **kwargs):
        params = inspect.signature(func).parameters
        print(f"Parameters for function '{func.__name__}':")
        # Loop through the parameters and try to print their values
        for i, (name, param) in enumerate(params.items()):
            # If it's a positional argument
            if i < len(args):
                print(f"- {name} = {args[i]}")
            # If it's a keyword argument
            elif name in kwargs:
                print(f"- {name} = {kwargs[name]}")
            # If it's not provided, use default if exists
            elif param.default != param.empty:
                print(f"- {name} = {param.default}")
            else:
                print(f"- {name} has no value!")
        return func(*args, **kwargs)
    return wrapper

def read_smap(file_path):
    """ Function parses smap filed

    Args:
        file_path (str): relative path to smap file to be parsed

    Returns:
        df (pd.DataFrame) : DataFrame of SV calls extracted from smap file
    """
    with open(file_path, 'r') as file:
        lines = file.readlines()
    start_index = None
    for i, line in enumerate(lines):
        if line.startswith("#h"):
            start_index = i
            break
    if start_index is not None:
        header = lines[start_index].strip()  # Store the header line
        data = lines[start_index + 2:]  # Skip the header line and the following line
        df = pd.DataFrame([line.strip().strip().split('\t') for line in data],columns=header.split()[1:])  # Set column names excluding the '#h' prefix
        print(f"smap DataFrame shape : {df.shape[0]} rows x {df.shape[1]} columns")
        return df
    else:
        print('Unable to parse smap!')
        return None  # Return None if no header line is found

def read_aneuploidy(file_path):
    """Function parses aneuploidy calls

    Args:
        file_path (str): relative path to aneuploidy file to be parsed

    Returns:
        df (pd.DataFrame): DataFrame of Aneuploidy calls extracted from aneuploidy file
    """
    with open(file_path, 'r') as file:
        lines = file.readlines()
    start_index = None
    for i, line in enumerate(lines):
        if line.startswith("#chr"):
            start_index = i
            break
    if start_index is not None:
        header = lines[start_index].strip()  # Store the header line
        data = lines[start_index + 1:]  # Skip only the header line
        df = pd.DataFrame([line.strip().split() for line in data],columns=header.split()).rename(columns={'#chr':'chr'}) # Set column names
        print(f"Aneuploidy DataFrame shape : {df.shape[0]} rows x {df.shape[1]} columns\n")
        return df
    else:
        print('Unable to parse Aneuploidy file!\n')
        return None 

def read_cnv(file_path):
    """Function parses CNV calls

    Args:
        file_path (str): relative path to CNV file to be parsed

    Returns:
        df (pd.DataFrame): DataFrame of CNV calls extracted from CNV file
    """
    with open(file_path, 'r') as file:
        lines = file.readlines()
    start_index = None
    for i, line in enumerate(lines):
        if line.startswith("#Id"):
            start_index = i
            break
    if start_index is not None:
        header = lines[start_index].strip()  # Store the header line
        data = lines[start_index + 1:]  # Skip only the header line
        df = pd.DataFrame([line.strip().split() for line in data],columns=header.split()).rename(columns={'#Id':'Id','Chromosome':'chr'}) # Set column names
        df['Start'] = df['Start'].astype(int)
        df['End'] = df['End'].astype(int)
        print(f"CNV DataFrame shape : {df.shape[0]} rows x {df.shape[1]} columns")
        return df
    else:
        print('Unable to parse CNV file!')
        return None 

def calc_difference(cnv):
    """Calculate the difference between start and end positions of subsequent rows based on chromosome and SVtype

    Args:
        cnv (pd.DataFrame): DataFrame consisting of CNV calls

    Returns:
        pd.Series: base pair difference between CNV segments on the same chromosome
    """
    vals = (cnv['Start'] - cnv['End'].shift(1)).fillna(cnv.groupby(['chr', 'Type'])['Start'].transform('first'))
    return vals

def stitching(cnv, cnv_stitch_window=550000, width=True):
    """Stitchs CNV calls within a particular window bin together

    Args:
        cnv (pd.DataFrame): DataFrame of CNV calls extracted from CNV file parsed by read_cnv()
        cnv_stitch_window (int, optional): Designated window to collapse CNV calls by. Defaults to 550000.
        width (bool, optional): Return the width of the CNV segment. Defaults to True.

    Returns:
        pd.DataFrame: CNV stitched calls
    """
    print(f"Stiching CNV DataFrame shape: {cnv.shape[0]} rows x {cnv.shape[1]} columns")
    print(f'Collapsing CNV calls using cnv stitch window : {cnv_stitch_window} (bp)')
    cnv = cnv.sort_values(by=['chr', 'Type', 'Start'])
    if cnv.shape[0] == 0:
        stitched_frame = cnv
        stitched_frame['diff'] = None
    else:
        cnv['diff'] = cnv.groupby(['chr', 'Type']).apply(calc_difference).T.values
        stitched_data = []
        cnv_calls = cnv.groupby(['chr', 'Type'])
        for (Chr, Type), cnv_frame in cnv_calls:
            stitched_cnv = collapse_rows(cnv_frame, cnv_stitch_window=cnv_stitch_window)
            stitched_data.append(stitched_cnv)
        stitched_frame = pd.concat(stitched_data)
    stitched_frame['Width'] = stitched_frame['End'] - stitched_frame['Start']
    return stitched_frame


def collapse_rows(df, cnv_stitch_window):
    """Collapses rows in a DataFrame based on a condition
    Args:
        df (pd.DataFrame): Input DataFrame
        cnv_stitch_window (int): Threshold value for collapsing rows
    Returns:
        pd.DataFrame: Collapsed DataFrame
    """
    collapsed_data = []
    current_start = None
    current_end = None
    current_row = None
    for index, row in df.iterrows():
        if current_start is None:
            current_start = row['Start']
            current_end = row['End']
            current_row = row
        else:
            diff = row['diff']
            if diff < cnv_stitch_window:
                current_end = row['End']
                current_row['End'] = current_end  # Update the 'End' value in the current row
            else:
                collapsed_data.append(current_row)
                current_start = row['Start']
                current_end = row['End']
                current_row = row
    if current_row is not None:
        collapsed_data.append(current_row)
    collapsed_df = pd.DataFrame(collapsed_data)
    return collapsed_df

def calculate_overlap_percentage(xlist,ylist):
    """
    Calculates the overlap percentage of two sets of values.
    Arguments:
    xlist -- First set of values as a tuple (start1, end1).
    ylist -- Second set of values as a tuple (start2, end2).
    Returns:
    overlap_percentage -- The overlap percentage as a float.
    """
    min1 = min(xlist)
    max1 = max(xlist)
    min2 = min(ylist)
    max2 = max(ylist)
    overlap = max(0, min(max1, max2) - max(min1, min2))
    length = max1-min1 + max2-min2
    return 2*overlap/length

def calculate_overlap_percentage_by_row(df, threshold=0.6, size_threshold=0.1):
    """
    Calculates the overlap percentage of two sets of values for every row relative to the first row.
    If all rows overlap by more than a given threshold and the average delta of 'Case Event Size' ratios
    is below size_threshold, return only the first row. Otherwise, return the first row and any rows 
    that don't meet these criteria. If 'Case Event Size' is NaN, assume the sizes are equal.
    Arguments:
    df -- pandas DataFrame containing columns "Case Event Start", "Case Event End", and "Case Event Size"
    threshold -- percentage threshold for filtering rows based on overlap (default 0.95)
    size_threshold -- threshold for filtering rows based on 'Case Event Size' ratios (default 0.1)
    Returns:
    df_filtered -- Filtered DataFrame.
    """
    df = df.sort_values(['Found in Control','Case Molecule Count'],ascending=[False,False])
    df.reset_index(inplace=True,drop=True)
    overlap_percentages = []
    size_ratios = []
    start1, end1, size1 = df.iloc[0][["Case Event Start", "Case Event End", "Case Event Size"]]
    min1, max1 = min(start1, end1), max(start1, end1)
    for i in range(len(df)):
        start2, end2, size2 = df.iloc[i][["Case Event Start", "Case Event End", "Case Event Size"]]
        min2, max2 = min(start2, end2), max(start2, end2)
        overlap = max(0, min(max1, max2) - max(min1, min2))
        length = (max1 - min1) + (max2 - min2)
        overlap_percentage = 2 * overlap / length if length != 0 else 0
        # If 'Case Event Size' is NaN, assume the sizes are equal
        if pd.isna(size1) or pd.isna(size2):
            size_ratio = 0
        else:
            size_ratio = abs(size1/size2 - 1) + abs(size2/size1 - 1)
        overlap_percentages.append(overlap_percentage)
        size_ratios.append(size_ratio)
    df['Overlap Percentage'] = overlap_percentages
    df['Size Ratio'] = size_ratios
    # If all rows (except the first) overlap by more than the threshold and the average size ratio is less than size_threshold
    # then return only the first row
    if all(percentage > threshold for percentage in overlap_percentages[1:]) and np.mean(size_ratios[1:]) < size_threshold:
        return df.iloc[[0]]
    # Otherwise, return the first row and any rows that don't meet these criteria
    else:
        return df[(df.index == 0) | (df['Overlap Percentage'] <= threshold) | (df['Size Ratio'] >= size_threshold)]

def pairwise_comparison(case, control, cnv_overlap_percentage=0.5, cnv_window=1000):
    """Performs pairwise comparison between case and control DataFrames
    Args:
        case (pd.DataFrame): DataFrame representing the case data
        control (pd.DataFrame): DataFrame representing the control data
        cnv_overlap_percentage (float, optional): Maximum allowed reciprocal overlap percentage. Defaults to 0.5.
        cnv_window (int, optional): Window size for Start and End comparison. Defaults to 1000.
    Returns:
        pd.DataFrame: Filtered case DataFrame
    """
    case['Width'] = case['Width'].astype(float)
    control['Width'] = control['Width'].astype(float)
    control['unique_id'] = control[['Start','End','Width','diff']].sum(axis=1).map(hash)
    case['unique_id'] = case[['Start','End','Width','diff']].sum(axis=1).map(hash)
    grouped_case = case.groupby('chr')
    grouped_control = control.groupby('chr')
    filtered_rows = []
    case_overlap = []
    case_checked = []
    control_checked = []
    print('\n--- Processing Case & Control CNV calls ---\n')
    for chr_name, case_group in grouped_case:
        if chr_name in grouped_control.groups:
            control_group = grouped_control.get_group(chr_name)
            print(f'Processing Chromosome : {chr_name}')
            print(f'\tNumber of Control events : {control_group.shape[0]}')
            print(f'\tNumber of Case events : {case_group.shape[0]}')
            for _, case_row in case_group.iterrows():
                case_specific = case_row.copy().to_frame().T
                case_specific.columns = case_specific.columns + '_Case'
                case_id = case_specific['unique_id_Case'].values[0]
                n = 0
                append = False
                for _, control_row in control_group.iterrows():
                    control_specific = control_row.copy().to_frame().T
                    control_specific.columns = control_specific.columns + '_Control'
                    control_id = control_specific['unique_id_Control'].values[0]
                    start_diff = abs(case_row['Start'] - control_row['Start'])
                    end_diff = abs(case_row['End'] - control_row['End'])
                    width_ratio = case_row['Width'] / control_row['Width']
                    overlap_percent = calculate_overlap_percentage(set([case_row['Start'],case_row['End']]),set([control_row['Start'],control_row['End']]))
                    if start_diff >= cnv_window and end_diff >= cnv_window and overlap_percent <= cnv_overlap_percentage and case_id not in case_checked and control_id not in control_checked:
                        filtered_rows.append(case_row.to_frame().T)
                        break  # Move to the next case row
                    if n == 0:
                        if (control_id not in control_checked) and (start_diff <= cnv_window or end_diff <= cnv_window or overlap_percent >= cnv_overlap_percentage):
                            case_control_overlap = pd.concat([case_specific.reset_index(drop=True),control_specific.reset_index(drop=True)],axis=1)
                            case_checked.append(case_id)
                            control_checked.append(control_id)
                            n +=1
                            append = True
                    else:
                        if (control_specific['End_Control'].values[0] > case_control_overlap['End_Control'].values[0]) and (control_specific['Start_Control'].values[0] < case_control_overlap['End_Case'].values[0]):
                            case_control_overlap['End_Control'] = control_specific['End_Control'].values[0]
                            control_checked.append(control_id)
                if append: # append merged overlapping case control
                    case_overlap.append(case_control_overlap)
        else:
            [filtered_rows.append(x[1].to_frame().T) for x in case_group.iterrows()]
    if len(filtered_rows) == 0:
        filtered_case = pd.DataFrame(filtered_rows,columns=case.columns)
    else:
        filtered_case = pd.concat(filtered_rows)
    if len(case_overlap) > 0:
        case_overlap_frame = pd.concat(case_overlap)
        case_overlap_frame['Width_Control'] = case_overlap_frame['End_Control'] - case_overlap_frame['Start_Control']
    else:
        case_overlap_frame = pd.DataFrame(columns=['Id_Case', 'chr_Case', 'Start_Case', 'End_Case', 'Width_Case', 'Type_Case', 'fractionalCopyNumber_Case', 'CopyNumber_Case', 'AlleleFreq_Case', 'Mask_overlap_fract_Case', 'diff_Case', 'unique_id_Case', 'Id_Control', 'chr_Control', 'Start_Control', 'End_Control', 'Width_Control', 'Type_Control', 'fractionalCopyNumber_Control', 'CopyNumber_Control', 'AlleleFreq_Control', 'Mask_overlap_fract_Control', 'diff_Control', 'unique_id_Control'])
    return filtered_case, case_overlap_frame

def compare_aneuploidy_data(case, control, aneuploidy_overlap_percentage):
    """Compares case and control DataFrames and returns rows unique to case based on specified conditions
    
    Args:
        case (pd.DataFrame): DataFrame representing the case data
        control (pd.DataFrame): DataFrame representing the control data
        aneuploidy_overlap_percentage (float, optional): Maximum allowed overlap percentage difference. Defaults to 0.5.
    
    Returns:
        pd.DataFrame: Rows unique to case DataFrame
    """
    case['fractChrLen'] = case['fractChrLen'].astype(float)
    control['fractChrLen'] = control['fractChrLen'].astype(float)
    case['fractCN'] = case['fractCN'].astype(float)
    control['fractCN'] = control['fractCN'].astype(float)
    merged = case.merge(control, on=['chr', 'types'], suffixes=('_case', '_control'), how='left')
    merged['Found_in_control_sample_assembly'] = np.nan
    print(f'Number of Aneuploidy calls to compare : {merged.shape[0]}')
    filtered_rows = []
    merged_rows = []
    for _, row in merged.iterrows():
        if pd.isnull(row['fractChrLen_control']) or pd.isnull(row['fractCN_control']):
            filtered_rows.append(row)
            merged_rows.append(row)
        else:
            fractChrLen_diff = abs(row['fractChrLen_case'] - row['fractChrLen_control'])
            fractCN_diff = abs(row['fractCN_case'] - row['fractCN_control'])
            if fractChrLen_diff > aneuploidy_overlap_percentage and fractCN_diff > aneuploidy_overlap_percentage:
                filtered_rows.append(row)
                merged_rows.append(row)
            else:
                row['Found_in_control_sample_assembly'] = 'yes'
                merged_rows.append(row)
    merged_case = pd.DataFrame(merged_rows,columns=merged.columns)
    filtered_case = pd.DataFrame(filtered_rows,columns=merged.columns).iloc[:,:5]
    filtered_case.columns = case.columns[:5]
    return filtered_case, merged_case

def process_cnvs(case, control, cnv_overlap_percentage=0.5, cnv_window=1000):
    """Processes CNV calls and reports unique and paired case & control CNV calls
    Args:
        case (pd.DataFrame): DataFrame representing the case data
        control (pd.DataFrame): DataFrame representing the control data
        cnv_overlap_percentage (float, optional): Maximum allowed reciprocal overlap percentage. Defaults to 0.5.
        cnv_window (int, optional): Window size for Start and End comparison. Defaults to 1000.
    Returns:
        pd.DataFrame: Filtered case DataFrame
    """
    cnv_reindex_cols = ['Cell type', 'Type', 'Id', 'Unique to case', 'chr', 'RefcontigID2', 'Start', 'End', 'Width', 'AlleleFreq', 'Case ID', 'Found in case', 'Self_molecule_count', 'Control ID', 'Found_in_control_sample_molecules', 'Control_molecule_count', 'fractionalCopyNumber']
    cnv_reindex_cols_unique_to_comp = ['Cell type', 'Event Type', 'Id', 'Unique to case', 'ChrA location','ChrB location', 'Start', 'Stop', 'SV size', 'VAF', 'Case ID','Found in case', 'Case count', 'Control ID','Found_in_control_sample_molecules', 'Control count','fractionalCopyNumber_Case', 'fractionalCopyNumber_Control','Start_Control', 'Stop_Control', 'SV size Control', 'VAF Control','CallType', 'SV size control', 'Found_in_control_sample_assembly']    
    unique_case_cnvs, cnv_comp_table = pairwise_comparison(case=case,control=control, cnv_overlap_percentage=cnv_overlap_percentage, cnv_window=cnv_window)
    cnv_reindex_comp_cols = ['Cell type', 'Type_Case', 'Id_Case', 'Unique to case', 'chr_Case', 'RefcontigID2', 'Start_Case', 'End_Case', 'Width_Case', 'AlleleFreq_Case', 'Case ID', 'Found in case', 'Self_molecule_count', 'Control ID', 'Found_in_control_sample_molecules', 'Control_molecule_count', 'fractionalCopyNumber_Case', 'fractionalCopyNumber_Control','Start_Control', 'End_Control', 'Width_Control', 'AlleleFreq_Control']
    cnv_comp_table_indexed = cnv_comp_table.reindex(cnv_reindex_comp_cols, axis=1).rename(columns={'Type_Case':'Event Type', 'Id_Case':'Id','chr_Case':'ChrA location', 'RefcontigID2':'ChrB location','Start_Case':'Start','End_Case':'Stop', 'Width_Case':'SV size', 'Self_molecule_count':'Case count','Control_molecule_count':'Control count', 'AlleleFreq_Case':'VAF','AlleleFreq_Control':'VAF Control','Width_Control':'SV size Control','End_Control':'Stop_Control'})
    cnv_comp_table_indexed['CallType'] = 'CNV'
    cnv_comp_table_indexed['SV size control'] = cnv_comp_table_indexed['Stop_Control'] - cnv_comp_table_indexed['Start_Control']
    cnv_comp_table_indexed['Found_in_control_sample_assembly'] = 'yes'
    cnv_comp_table_indexed['Found in case'] = 'yes'
    unique_case_cnvs_subset = unique_case_cnvs.reindex(cnv_reindex_cols, axis=1).rename(columns={'Type':'Event Type', 'chr':'ChrA location', 'RefcontigID2':'ChrB location','RefStartPos':'Start','End':'Stop', 'Width':'SV size', 'Self_molecule_count':'Case count','Control_molecule_count':'Control count', 'AlleleFreq':'VAF'})
    unique_case_cnvs_subset['CallType'] = 'CNV'
    unique_case_to_merge = unique_case_cnvs_subset.rename(columns={'fractionalCopyNumber':'fractionalCopyNumber_Case'}).reindex(cnv_reindex_cols_unique_to_comp,axis=1)
    unique_case_to_merge['Found_in_control_sample_assembly'] = 'no'
    joined_cnv_comp = pd.concat([cnv_comp_table_indexed, unique_case_to_merge])
    return unique_case_cnvs_subset, joined_cnv_comp

def process_aneuploidies(case, control, aneuploidy_overlap_percentage):
    """Processes Aneuploidy calls and reports unique and paired case & control Aneuploidy calls
    
    Args:
        case (pd.DataFrame): DataFrame representing the case data
        control (pd.DataFrame): DataFrame representing the control data
        aneuploidy_overlap_percentage (float, optional): Maximum allowed overlap percentage difference. Defaults to 0.5.
    
    Returns:
        unique_case_aneu_subset (pd.DataFrame): Rows unique to case DataFrame
        aneu_comp_table_indexed (pd.DataFrame): Aneuploidy events shared across control and case
    """
    contig_length = {1: 248956422.0, 2: 242193529.0, 3: 198295559.0, 4: 190214555.0, 5: 181538259.0, 6: 170805979.0, 7: 159345973.0, 8: 145138636.0, 9: 138394717.0, 10: 133797422.0, 11: 135086622.0, 12: 133275309.0, 13: 114364328.0, 14: 107043718.0, 15: 101991189.0, 16: 90338345.0, 17: 83257441.0, 18: 80373285.0, 19: 58617616.0, 20: 64444167.0, 21: 46709983.0, 22: 50818468.0, 23: 156040895.0, 24: 57227415.0}
    aneu_reindex_cols = ['Cell type', 'types', 'Id', 'Unique to case', 'chr', 'RefcontigID2', 'Start', 'End', 'fractChrLen_case', 'AlleleFreq', 'Case ID', 'Found in case', 'Self_molecule_count', 'Control ID', 'Found_in_control_sample_assembly', 'Found_in_control_sample_molecules', 'Control_molecule_count', 'fractCN', 'SV size Control','fractionalCopyNumber_control','score_control','score']
    if (case.shape[0] != 0):
        unique_case_aneuploidies, aneu_comp_table = compare_aneuploidy_data(case=case, control=control, aneuploidy_overlap_percentage=aneuploidy_overlap_percentage)
        aneu_reindex_comp_cols = ['Cell type', 'types', 'Id', 'Unique to case', 'chr', 'RefcontigID2', 'Start', 'End', 'fractChrLen_case', 'AlleleFreq', 'Case ID', 'Found in case', 'Self_molecule_count', 'Control ID', 'Found_in_control_sample_assembly', 'Found_in_control_sample_molecules', 'Control_molecule_count', 'fractCN_case', 'fractChrLen_control', 'fractCN_control', 'score_control','score_case']
        aneu_comp_table_indexed = aneu_comp_table.reindex(aneu_reindex_comp_cols, axis=1).rename(columns={'types':'Event Type', 'chr':'ChrA location', 'RefcontigID2':'ChrB location','RefStartPos':'Start','End':'Stop','Self_molecule_count':'Case count','Control_molecule_count':'Control count', 'AlleleFreq':'VAF','fractCN_case':'fractionalCopyNumber','fractCN_control':'fractionalCopyNumber_control'})
        aneu_comp_table_indexed['CallType'] = 'Aneuploidy'
        aneu_comp_table_indexed['fractChrLen_case'] = aneu_comp_table_indexed.apply(lambda row: row['fractChrLen_case'] * contig_length[int(row['ChrA location'])], axis=1)
        aneu_comp_table_indexed['fractChrLen_control'] = aneu_comp_table_indexed.apply(lambda row: row['fractChrLen_control'] * contig_length[int(row['ChrA location'])], axis=1)
        aneu_comp_table_indexed['Found in case'] = 'yes'
        aneu_comp_table_indexed = aneu_comp_table_indexed.rename(columns={'fractChrLen_case':'SV size','fractChrLen_control':'SV size Control'})
        unique_case_aneu_subset = unique_case_aneuploidies.reindex(aneu_reindex_cols, axis=1).rename(columns={'types':'Event Type', 'chr':'ChrA location', 'RefcontigID2':'ChrB location','RefStartPos':'Start','End':'Stop', 'fractChrLen_case':'SV size', 'Self_molecule_count':'Case count','Control_molecule_count':'Control count', 'AlleleFreq':'VAF','fractCN':'fractionalCopyNumber','score':'score_case'})
        unique_case_aneu_subset['CallType'] = 'Aneuploidy'  
        unique_case_aneu_subset['ChrB location'] = unique_case_aneu_subset['ChrA location']
    else:
        empty_cols = ['Cell type', 'Event Type', 'Id', 'Unique to case', 'ChrA location', 'ChrB location', 'Start', 'Stop', 'SV size', 'VAF', 'Case ID', 'Found in case', 'Case count', 'Control ID', 'Found_in_control_sample_assembly','Found_in_control_sample_molecules', 'Control count', 'fractionalCopyNumber', 'SV size control', 'fractionalCopyNumber_control', 'CallType']
        aneu_comp_table_indexed = pd.DataFrame(columns=empty_cols)
        unique_case_aneu_subset = pd.DataFrame(columns=aneu_reindex_cols)  
    return unique_case_aneu_subset, aneu_comp_table_indexed

def process_smap(dual_smap_frame):
    """Processes SV calls and reports unique and paired case & control SV calls

    Args:
        dual_smap_frame (pd.DataFrame): DataFrame consisting of dual annotation SV calls
    """
    smap_reindex_cols = ['Cell type', 'Type', 'SmapEntryID', 'Unique to case', 'RefcontigID1', 'RefcontigID2', 'RefStartPos', 'RefEndPos', 'SVsize', 'VAF', 'Case ID', 'Found in case', 'Self_molecule_count', 'Control ID', 'Control_molecule_count', 'fractionalCopyNumber','Found_in_control_sample_molecules','Found_in_control_sample_assembly']
    smap_subset = dual_smap_frame.reindex(smap_reindex_cols, axis=1).rename(columns={'SmapEntryID':'Id', 'Type':'Event Type', 'RefcontigID1':'ChrA location', 'RefcontigID2':'ChrB location','RefStartPos':'Start','RefEndPos':'Stop', 'SVsize':'SV size', 'Self_molecule_count':'Case count','Control_molecule_count':'Control count'})
    smap_subset['CallType'] = 'SV'
    smap_subset['Found in case'] = 'yes'
    smap_filtered = smap_subset[(smap_subset['Found_in_control_sample_molecules'] != 'yes') & (smap_subset['Found_in_control_sample_assembly'] != 'yes')]

    return smap_filtered, smap_subset

def convert_to_int(x):
    try:
        return int(float(x))
    except (ValueError, TypeError):
        return None

def process_all_calls(smap_subset, aneu_comp_table_indexed, cnv_comp_table_indexed, control_smap_frame):
    """ Aggregates and formats all SV, CNV and Aneuploidy calls into dataframe

    Args:
        smap_subset (pd.DataFrame): smap frame returned by read_smap
        aneu_comp_table_indexed (pd.DataFrame): aneuploidy frame returned by process_aneuploidies
        cnv_comp_table_indexed (pd.DataFrame): CNV frame returned by process_cnvs
        control_smap_frame (pd.DataFrame): control smap frame
    Returns:
        all_calls (pd.DataFrame): reindexed frame containing all SV, CNV and Aneuploidy calls with accompanying control calls when available
    """
    reindex_cols = ['Cell type', 'Event Type', 'Id', 'ChrA location', 'ChrB location', 'Start','Start_Control', 'Stop', 'Stop_Control','SV size','SV size control','VAF', 'VAF Control', 'Case count', 'Control count','fractionalCopyNumber_Case', 'fractionalCopyNumber_Control', 'Case ID', 'Control ID', 'Unique to case', 'Found in case', 'Found_in_control_sample_molecules', 'Found_in_control_sample_assembly', 'CallType']
    reindex_final_cols = ['Cell type', 'Event Type', 'CallType', 'variant_id_case', 'chrA_case', 'chrB_case', 'start_case', 'end_case', 'size_case', 'VAF_case', 'start_control', 'end_control', 'size_control', 'VAF_control', 'fractional_copy_number_case', 'fractional_copy_number_control', 'molecule_count_case', 'molecule_count_control', 'found_in_control_molecules', 'Found_in_control_sample_assembly', 'Found in control', 'Unique to case', 'ID_case', 'ID_control']
    map_dict = {'ChrA location':'chrA_case','ChrB location':'chrB_case','Start':'start_case','Stop':'end_case','SV size':'size_case','VAF':'VAF_case','Start_Control':'start_control','Stop_Control':'end_control','SV size control':'size_control','VAF Control':'VAF_control','fractionalCopyNumber_Case':'fractional_copy_number_case','fractionalCopyNumber_Control':'fractional_copy_number_control','Case count':'molecule_count_case','Control count':'molecule_count_control','Found_in_control_sample_molecules':'found_in_control_molecules','Id':'variant_id_case'}   
    smap_reindexed = smap_subset.reindex(reindex_cols,axis=1)
    smap_subset = associate_sv_with_controls(smap_reindexed, control_smap_frame)
    aneu_comp_table_indexed = aneu_comp_table_indexed.rename(columns={'fractionalCopyNumber':'fractionalCopyNumber_Case','fractionalCopyNumber_control':'fractionalCopyNumber_Control','SV size Control':'SV size control'}).reindex(reindex_cols,axis=1)
    cnv_comp_table_indexed = cnv_comp_table_indexed.reindex(reindex_cols,axis=1)
    smap_subset.reset_index(drop=True,inplace=True)
    aneu_comp_table_indexed.reset_index(drop=True,inplace=True)
    cnv_comp_table_indexed.reset_index(drop=True,inplace=True)
    all_calls = pd.concat([smap_subset, aneu_comp_table_indexed, cnv_comp_table_indexed],ignore_index=True,join='outer')
    all_calls = all_calls.rename(columns=map_dict)
    all_calls['VAF_case'] = all_calls['VAF_case'].astype(float)
    all_calls['VAF_control'] = all_calls['VAF_control'].astype(float)
    all_calls['fractional_copy_number_case'] = all_calls['fractional_copy_number_case'].astype(float)
    all_calls['fractional_copy_number_control'] = all_calls['fractional_copy_number_control'].astype(float)
    all_calls['size_case'] = all_calls['size_case'].apply(convert_to_int).astype('Int64')
    all_calls['size_control'] = all_calls['size_control'].apply(convert_to_int).astype('Int64')
    all_calls['start_case'] = all_calls['start_case'].apply(convert_to_int).astype('Int64')
    all_calls['end_case'] = all_calls['end_case'].apply(convert_to_int).astype('Int64')
    all_calls['start_control'] = all_calls['start_control'].apply(convert_to_int).astype('Int64')
    all_calls['end_control'] = all_calls['end_control'].apply(convert_to_int).astype('Int64')
    all_calls = all_calls.reindex(reindex_final_cols,axis=1)
    all_calls['Found in control'] = all_calls.apply(lambda row: 'yes' if row['found_in_control_molecules'] == 'yes' or row['Found_in_control_sample_assembly'] == 'yes' else 'no', axis=1)
    return all_calls

def process_subset_calls(out_table):
    """ Aggregates and formats all SV, CNV and Aneuploidy calls into dataframe

    Args:
        out_table (pd.DataFrame): merged CNV, aneuploidy and SV frame unique to case 
    Returns:
        out_table_reindexed (pd.DataFrame): reindexed frame containing unique case specific SV, CNV and Aneuploidy calls
    """
    reindex_cols = ['Cell type', 'Event Type', 'Id', 'ChrA location', 'ChrB location', 'Start','Start_Control', 'Stop', 'Stop_Control','SV size','SV size control','VAF', 'VAF Control', 'Case count', 'Control count','fractionalCopyNumber', 'fractionalCopyNumber_control', 'Case ID', 'Control ID', 'Unique to case', 'Found in case', 'Found_in_control_sample_molecules', 'Found_in_control_sample_assembly', 'CallType']
    reindex_final_cols = ['Cell type', 'ID_case', 'ID_control', 'CallType','Event Type','variant_id_case', 'chrA_case', 'chrB_case', 'start_case', 'end_case', 'size_case', 'VAF_case', 'start_control', 'end_control', 'size_control', 'VAF_control', 'fractional_copy_number_case', 'fractional_copy_number_control', 'molecule_count_case', 'molecule_count_control', 'found_in_control_molecules', 'Found_in_control_sample_assembly', 'Unique to case']
    map_dict = {'ChrA location':'chrA_case','ChrB location':'chrB_case','Start':'start_case','Stop':'end_case','SV size':'size_case','VAF':'VAF_case','Start_Control':'start_control','Stop_Control':'end_control','SV size control':'size_control','VAF Control':'VAF_control','fractionalCopyNumber':'fractional_copy_number_case','fractionalCopyNumber_control':'fractional_copy_number_control','Case count':'molecule_count_case','Control count':'molecule_count_control','Found_in_control_sample_molecules':'found_in_control_molecules','Id':'variant_id_case'}
    out_table_reindexed = out_table.reindex(reindex_cols,axis=1)
    out_table_reindexed = out_table_reindexed.rename(columns=map_dict)
    out_table_reindexed['VAF_case'] = out_table_reindexed['VAF_case'].astype(float)
    out_table_reindexed['VAF_control'] = out_table_reindexed['VAF_control'].astype(float)
    out_table_reindexed['size_case'] = out_table_reindexed['size_case'].apply(convert_to_int).astype('Int64')
    out_table_reindexed['size_control'] = out_table_reindexed['size_control'].apply(convert_to_int).astype('Int64')
    out_table_reindexed['start_case'] = out_table_reindexed['start_case'].apply(convert_to_int).astype('Int64')
    out_table_reindexed['end_case'] = out_table_reindexed['end_case'].apply(convert_to_int).astype('Int64')
    out_table_reindexed['start_control'] = out_table_reindexed['start_control'].apply(convert_to_int).astype('Int64')
    out_table_reindexed['end_control'] = out_table_reindexed['end_control'].apply(convert_to_int).astype('Int64')
    out_table_reindexed['size_case'] = out_table_reindexed['size_case'].replace(-1,np.nan)
    out_table_reindexed['size_control'] = out_table_reindexed['size_control'].replace(-1,np.nan)
    out_table_reindexed = out_table_reindexed.reindex(reindex_final_cols,axis=1)
    return out_table_reindexed

def associate_sv_with_controls(smap_subset, control_smap_frame):
    """ Aligns control SVs with Case using dual VAP default parameters
    Args:
        smap_subset (pd.DataFrame): _description_
        control_smap_frame (pd.DataFrame): _description_
    """
    ins_del_position_overlap = 10000
    ins_del_size_percent_similarity = 50
    inversion_position_overlap = 50000
    translocation_position_overlap = 50000
    duplication_position_overlap = 10000
    duplication_size_percent_similarity = 50
    reindex_cols = ['Cell type', 'Event Type', 'Id', 'ChrA location', 'ChrB location', 'Start','Start_Control', 'Stop', 'Stop_Control','SV size','SV size control','VAF', 'VAF Control', 'Case count', 'Control count','fractionalCopyNumber', 'fractionalCopyNumber_control', 'Case ID', 'Control ID', 'Unique to case', 'Found in case', 'Found_in_control_sample_molecules', 'Found_in_control_sample_assembly', 'CallType']
    smap_subset = smap_subset.reindex(reindex_cols,axis=1)
    smap_subset['Start_Control'] = np.nan
    smap_subset['Stop_Control'] = np.nan
    smap_subset['VAF Control'] = np.nan
    smap_subset['Start'] = smap_subset['Start'].astype(float)
    smap_subset['Stop'] = smap_subset['Stop'].astype(float)
    control_smap_frame['RefStartPos'] = control_smap_frame['RefStartPos'].astype(float)
    control_smap_frame['RefEndPos'] = control_smap_frame['RefEndPos'].astype(float)
    smap_subset_updated = smap_subset.copy()
    print('\n--- Intersecting SV Case & Control calls ---\n')
    grouped_calls = smap_subset.groupby(['Event Type','ChrA location'])
    for (event_type, chr1), smap_subset_frame in grouped_calls:
        print(f'Processing SVType : {event_type}\n\tChromosome : {chr1}')
        print(f'\tNumber of events : {smap_subset_frame.shape[0]}\n')
        control_subset = control_smap_frame[(control_smap_frame['Type'] == event_type) & (control_smap_frame['RefcontigID1'] == chr1)]
        if (event_type == 'insertion') or (event_type == 'deletion'):
            for _, row in smap_subset_frame.iterrows():
                sub_control_frame = control_subset[(control_subset['RefStartPos'].between(row['Start']-ins_del_position_overlap,row['Start']+ins_del_position_overlap)) & (control_subset['RefEndPos'].between(row['Stop']-ins_del_position_overlap,row['Stop']+ins_del_position_overlap))]
                overlap_percent = [(float(calculate_overlap_percentage(set([row['Start'],row['Stop']]),set([control_row['RefStartPos'],control_row['RefEndPos']])))*100.0) > ins_del_size_percent_similarity for _,control_row in sub_control_frame.iterrows()]
                sub_control_frame = sub_control_frame.loc[overlap_percent]
                if sub_control_frame.shape[0] > 0:
                    sub_control_values = sub_control_frame.loc[:,['RefStartPos','RefEndPos', 'VAF', 'SVsize']].iloc[0,:].to_frame().T
                    smap_subset_updated.loc[row.name,'VAF Control'] = sub_control_values['VAF'].values[0]
                    smap_subset_updated.loc[row.name,'Start_Control'] = sub_control_values['RefStartPos'].values[0]
                    smap_subset_updated.loc[row.name,'Stop_Control'] = sub_control_values['RefEndPos'].values[0]
                    smap_subset_updated.loc[row.name,'SV size control'] = sub_control_values['SVsize'].values[0]
        if (event_type.startswith('inversion')):
            for _, row in smap_subset_frame.iterrows():
                sub_control_frame = control_subset[(control_subset['RefStartPos'].between(row['Start']-inversion_position_overlap,row['Start']+inversion_position_overlap)) & (control_subset['RefEndPos'].between(row['Stop']-inversion_position_overlap,row['Stop']+inversion_position_overlap))]
                if sub_control_frame.shape[0] > 0:
                    sub_control_values = sub_control_frame.loc[:,['RefStartPos','RefEndPos', 'VAF', 'SVsize']].iloc[0,:].to_frame().T
                    smap_subset_updated.loc[row.name,'VAF Control'] = sub_control_values['VAF'].values[0]
                    smap_subset_updated.loc[row.name,'Start_Control'] = sub_control_values['RefStartPos'].values[0]
                    smap_subset_updated.loc[row.name,'Stop_Control'] = sub_control_values['RefEndPos'].values[0]
                    smap_subset_updated.loc[row.name,'SV size control'] = sub_control_values['SVsize'].values[0]  
        if (event_type.startswith('trans')):
            for _, row in smap_subset_frame.iterrows():
                sub_control_frame = control_subset[(control_subset['RefStartPos'].between(row['Start']-translocation_position_overlap,row['Start']+translocation_position_overlap)) & (control_subset['RefEndPos'].between(row['Stop']-translocation_position_overlap,row['Stop']+translocation_position_overlap))]
                if sub_control_frame.shape[0] > 0:
                    sub_control_values = sub_control_frame.loc[:,['RefStartPos','RefEndPos', 'VAF', 'SVsize']].iloc[0,:].to_frame().T
                    smap_subset_updated.loc[row.name,'VAF Control'] = sub_control_values['VAF'].values[0]
                    smap_subset_updated.loc[row.name,'Start_Control'] = sub_control_values['RefStartPos'].values[0]
                    smap_subset_updated.loc[row.name,'Stop_Control'] = sub_control_values['RefEndPos'].values[0]
                    smap_subset_updated.loc[row.name,'SV size control'] = sub_control_values['SVsize'].values[0] 
        if (event_type.startswith('duplication')):
            for _, row in smap_subset_frame.iterrows():
                sub_control_frame = control_subset[(control_subset['RefStartPos'].between(row['Start']-duplication_position_overlap,row['Start']+duplication_position_overlap)) & (control_subset['RefEndPos'].between(row['Stop']-duplication_position_overlap,row['Stop']+duplication_position_overlap))]
                overlap_percent = [(float(calculate_overlap_percentage(set([row['Start'],row['Stop']]),set([control_row['RefStartPos'],control_row['RefEndPos']])))*100.0) > duplication_size_percent_similarity for _,control_row in sub_control_frame.iterrows()]
                sub_control_frame = sub_control_frame.loc[overlap_percent]
                if sub_control_frame.shape[0] > 0:
                    sub_control_values = sub_control_frame.loc[:,['RefStartPos','RefEndPos', 'VAF', 'SVsize']].iloc[0,:].to_frame().T
                    smap_subset_updated.loc[row.name,'VAF Control'] = sub_control_values['VAF'].values[0]
                    smap_subset_updated.loc[row.name,'Start_Control'] = sub_control_values['RefStartPos'].values[0]
                    smap_subset_updated.loc[row.name,'Stop_Control'] = sub_control_values['RefEndPos'].values[0]
                    smap_subset_updated.loc[row.name,'SV size control'] = sub_control_values['SVsize'].values[0]
    return smap_subset_updated

def reorder_sheet(all_calls):
    """ Function reorders and subsets results based on desired columns and split data by SV and CNV/Aneuploidy calls

    Args:
        all_calls (pd.DataFrame): All intersected calls (SV/CNV) (Case/Control)

    Returns:
        cnv_calls_final (pd.DataFrame): Reordered, final CNV calls
        sv_calls_final (pd.DataFrame): Reordered, final SV calls
    """
    reindex_final_cols = ['Cell type', 'ID_case', 'ID_control', 'CallType', 'Event Type', 'chrA_case', 'chrB_case', 'start_case', 'end_case', 'size_case', 'VAF_case', 'start_control', 'end_control', 'size_control', 'VAF_control', 'fractional_copy_number_case', 'fractional_copy_number_control', 'molecule_count_case', 'molecule_count_control', 'Found_in_control_sample_assembly', 'found_in_control_molecules', 'Found in control']
    reindex_final_sv_cols = ['Cell type', 'ID_case', 'ID_control', 'Event Type', 'chrA_case', 'chrB_case', 'start_case', 'end_case', 'size_case','molecule_count_case', 'molecule_count_control','Found in control']
    reindex_final_cnv_cols = ['Cell type', 'ID_case', 'ID_control', 'CallType', 'Event Type', 'chrA_case', 'start_case', 'end_case', 'size_case','fractional_copy_number_case','fractional_copy_number_control','Found in control']
    name_sv_map_dict = {'Cell type':'Cell Type', 'ID_case': 'Case Sample Name','ID_control':'Control Sample Name','Event Type':'Event Type', 'chrA_case':'Case Start Chromosome','chrB_case':'Case End Chromosome','start_case':'Case Event Start','end_case':'Case Event End','size_case':'Case Event Size','molecule_count_case':'Case Molecule Count','molecule_count_control':'Control Molecule Count', 'Found in control':'Found in Control'}
    name_cnv_map_dict = {'Cell type':'Cell Type', 'ID_case': 'Case Sample Name','ID_control':'Control Sample Name','CallType':'Call Type','Event Type':'Event Type','chrA_case':'Case Chromosome','start_case':'Case Event Start','end_case':'Case Event End','size_case':'Case Event Size','fractional_copy_number_case':'Case Fractional Copy Number','fractional_copy_number_control':'Control Fractional Copy Number', 'Found in control':'Found in Control'}
    all_calls = all_calls.reindex(reindex_final_cols,axis=1)
    all_calls['size_case'] = all_calls['size_case'].replace(-1,np.nan)
    all_calls['size_control'] = all_calls['size_control'].replace(-1,np.nan)
    cnv_calls = all_calls[all_calls['CallType']!='SV']
    cnv_calls_final = cnv_calls.reindex(reindex_final_cnv_cols,axis=1).rename(columns=name_cnv_map_dict)
    sv_calls = all_calls[all_calls['CallType'] =='SV']
    sv_calls_final = sv_calls.reindex(reindex_final_sv_cols,axis=1).rename(columns=name_sv_map_dict)

    return cnv_calls_final, sv_calls_final

def process_json(input_json):
    """ Function parses json to extract CNV statistics and coerces dictionary into DataFrame

    Args:
        input_json (str): relative path to input json
    Returns
    """
    with open(input_json) as json_file:
        data = json.load(json_file)
    cnv_statistics = pd.DataFrame.from_dict(data['CNV Statistics']['value']).T.iloc[1:,:].loc[:,['value','description']]
    cnv_map_value = {'Percent above expected (2 Mbp window)':20, 'Percent above expected (6 Mbp window)':20,'Correlation with label density':0.25,'Wave template correlation':0.4}
    cnv_statistics['Metrics_Passed'] = cnv_statistics['value'].apply(lambda x: 'Fail' if cnv_statistics.index[cnv_statistics['value'] == x][0] in cnv_map_value and float(x) > cnv_map_value[cnv_statistics.index[cnv_statistics['value'] == x][0]] else ('Pass' if cnv_statistics.index[cnv_statistics['value'] == x][0] in cnv_map_value else 'NA'))
    cnv_statistics.index.name='Metric'
    cnv_indx = ['Percent above expected (2 Mbp window)','Percent above expected (6 Mbp window)','Correlation with label density','Wave template correlation']
    cnv_statistics_subset = cnv_statistics.loc[cnv_indx,:]
    return cnv_statistics_subset

def process_case_and_control_json(control_json, case_json, control_id, case_id):
    """ Function aggregates control and case json and returns a merged frame

    Args:
        control_json (str): relative path to control json
        case_json (str): relative path to case json
        control_id (str): control id
        case_id (str): case id
    """
    control_cnv = process_json(control_json)
    control_cnv.columns = [control_id, 'description', 'Control QC Passed']
    control_cnv[control_id] = control_cnv[control_id].astype(float)
    case_cnv = process_json(case_json)
    case_cnv.columns = [case_id, 'description', 'Case QC Passed']
    case_cnv[case_id] = case_cnv[case_id].astype(float)
    joined_cnv_calls = pd.concat([case_cnv.reindex([case_id,'Case QC Passed'],axis=1),control_cnv.reindex([control_id,'Control QC Passed'],axis=1)],axis=1)
    include_cnv_calls = (joined_cnv_calls['Control QC Passed'] == 'Fail').sum() + (joined_cnv_calls['Case QC Passed'] == 'Fail').sum()
    joined_cnv_calls.index = ['Percent difference between observed and expected coefficient of variation (2 Mbp window)', 'Percent difference between observed and expected coefficient of variation (6 Mbp window)', 'Correlation with label density', 'Wave template correlation']
    joined_cnv_calls.index.name = 'Metric'
    if include_cnv_calls == 0:
        print(f"\nCNV criteria passed for both Case and Control samples\n")
    else:
        print(f"\nCNV criteria failed.\n\tCase contains [{(joined_cnv_calls['Case QC Passed'] == 'Fail').sum()}] failed metrics\n\tControl contains [{(joined_cnv_calls['Control QC Passed'] == 'Fail').sum()}] failed metrics")
    return joined_cnv_calls, include_cnv_calls


def filter_duplicate_calls(sv_calls):
    """ Filters duplicate map calls made by RVP pipeline
    Args:
        sv_calls (pd.DataFrame): Frame returned from reorder_sheet()
    """
    print('\n--- Filtering duplicate SV calls ---')
    print(f'Input SV calls : {sv_calls.shape[0]}')
    sv_calls.reset_index(inplace=True,drop=True)
    sv_calls['Case Molecule Count'] = sv_calls['Case Molecule Count'].astype('int64')
    sv_calls['Control Molecule Count'] = sv_calls['Control Molecule Count'].astype('int64')
    sv_calls_sorted = sv_calls.sort_values(['Event Type','Case Start Chromosome','Found in Control','Case Molecule Count'],ascending=[True,True,False,False])
    grouped_calls = sv_calls_sorted.groupby(['Event Type','Case Start Chromosome'])
    final_calls = []
    for (event_type, chrom),subset_frame in grouped_calls:
        if subset_frame.shape[0] == 1:
            final_calls.append(subset_frame)
        else:
            subset_frame = calculate_overlap_percentage_by_row(subset_frame).iloc[:,:-2]
            final_calls.append(subset_frame)
    sv_final_calls = pd.concat(final_calls)
    print(f'Final SV calls : {sv_final_calls.shape[0]}')
    print(f'Filtered duplicate calls : {sv_calls.shape[0] - sv_final_calls.shape[0]}')
    return sv_final_calls

def format_rounded(number, decimals=3):
    """Formats a number to a fixed number of decimals, retaining trailing zeros.

    Args:
        number (float|int): The number to be formatted.
        decimals (int, optional): The number of decimals to format to. Defaults to 3.

    Returns:
        str: The formatted string representation of the number.
    """
    if pd.isna(number) or np.isnan(number):  # Handle NaN values
        return ''
    return f"{round(number, decimals):.{decimals}f}"

def write_dataframe_to_csv(df, filename, decimals=3):
    """Writes a DataFrame to a CSV file after transforming numeric columns 
    to strings with a fixed number of trailing zeros.

    Args:
        df (pd.DataFrame): The DataFrame to be written to CSV.
        filename (str): The name (or path) of the output CSV file.
        decimals (int, optional): The number of decimals to format to. Defaults to 3.
    """
    
    # Transform numeric columns to strings with trailing zeros
    for col in df.select_dtypes(include=['float64', 'int64']).columns:
        df[col] = df[col].apply(lambda x: format_rounded(x, decimals=decimals))
    
    # Write to CSV
    df.to_csv(filename, index=False, na_rep='NA')

def generate_docx(frame, file_handle):
    """ This function 

    Args:
        frame (pd.DataFrame): DataFrame to be converted to docx format
        file_handle (str): file handle
    """
    doc = docx.Document()
    t = doc.add_table(rows=frame.shape[0] + 1, cols=frame.shape[1])
    t.style = 'TableGrid'

    # Set up the desired font and alignment
    font_name = "Times New Roman"
    center_alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the column headings
    for j in range(frame.shape[1]):
        cell = t.cell(0, j)
        cell.paragraphs[0].clear()  # Clear the paragraph
        run = cell.paragraphs[0].add_run(frame.columns[j])  # Add a new run with the header text
        cell.paragraphs[0].alignment = center_alignment
        run.font.name = font_name
        run.font.size = Pt(12)
    for i in range(frame.shape[0]):
        for j in range(frame.shape[1]):
            value = frame.iat[i, j]
            # Format the value if it's a float
            if isinstance(value, (float, int)):
                formatted_value = format_rounded(value)
            else:
                formatted_value = str(value)
            cell = t.cell(i + 1, j)
            cell.text = formatted_value
            cell.paragraphs[0].alignment = center_alignment
            for run in cell.paragraphs[0].runs:
                run.font.name = font_name
                run.font.size = Pt(12)
    doc.save('{}.docx'.format(file_handle))

def center_excel_cells(writer):
    """ Function to center values within a cell

    Args:
        writer (pd.ExcelWriter): Excel Writer
    """
    workbook = writer.book
    for sheetname in writer.sheets:
        worksheet = writer.sheets[sheetname]
        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = cell.alignment.copy(horizontal='center', vertical='center')

def extract_path_from_handle(file_handle):
    """
    Extract the path from a file handle.
    
    Parameters:
    - file_handle (str): The file handle string which may contain a path.
    
    Returns:
    - str: The path if present, otherwise an empty string.
    """
    path, _ = os.path.split(file_handle)
    return path

@print_parameters_on_call
def compare_calls(dual_aneuploidy, dual_smap, dual_cnv, control_aneuploidy, control_cnv, out_file, case_id, control_id, celltype, control_smap, control_json, case_json, cnv_overlap_percentage=0.3, aneuploidy_overlap_percentage=0.5, cnv_window=1000, cnv_stitch_window=550000):
    """This function compares case and control Aneuploidy and CNV calls and reports case specific SV, CNV and Aneuploidy calls

    Args:
        dual_aneuploidy (str): relative path to aneuploidy file from dual annotation zip
        dual_smap (str): relative path to smap file from dual annotation zip
        dual_cnv (str): relative path to CNV file from dual annotation zip
        control_aneuploidy (str): relative path to aneuploidy file from control zip
        control_cnv (str): relative path to CNV file from control zip
        out_file (str): outfile handle
        case_id (str): case id
        control_id (str): control id
        celltype (str): cell type
        control_smap (str): relative path to smap file from control zip
        input_json (str): relative path to input json file that contains CNV metrics
        cnv_overlap_percentage (float): maximum reciprocal overlap percentage to consider CNV unique to case
        aneuploidy_overlap_percentag (float): maximum fractional coverage difference to consider Aneuploidies unique to case
        cnv_window (int): base pair window to extend start and stop positions by for CNV calls
        cnv_stitch_window (int): base pair window used to extend and join neighboring cnvs that fall within n-basepairs from one another. Defaults to 550000.
    """
    print('\n--- Processing smap files ---')
    print('Parsing dual annotation smap : {}'.format(dual_smap))
    dual_smap_frame = read_smap(dual_smap)

    print('\nParsing control smap : {}'.format(control_smap))
    control_smap_frame = read_smap(control_smap)

    print('\n--- Processing Case Aneuploidy data ---')
    dual_aneuploidy_frame = read_aneuploidy(dual_aneuploidy)

    print('--- Processing Case CNV data ---')
    dual_cnv_frame = read_cnv(dual_cnv)
    stitched_dual_cnvs = stitching(dual_cnv_frame, cnv_stitch_window=cnv_stitch_window)

    print('\n--- Processing Control Aneuploidy data ---')
    control_aneuploidy_frame = read_aneuploidy(control_aneuploidy)

    print('--- Processing Control CNV data ---')
    control_cnv_frame = read_cnv(control_cnv)
    stitched_control_cnvs = stitching(control_cnv_frame, cnv_stitch_window=cnv_stitch_window)

    unique_case_cnvs_subset, cnv_comp_table_indexed = process_cnvs(case=stitched_dual_cnvs, control=stitched_control_cnvs, cnv_overlap_percentage=cnv_overlap_percentage, cnv_window=cnv_window)
    print('\n--- Intersecting Aneuploidy calls ---')
    unique_case_aneu_subset, aneu_comp_table_indexed = process_aneuploidies(case=dual_aneuploidy_frame, control=control_aneuploidy_frame, aneuploidy_overlap_percentage=aneuploidy_overlap_percentage)

    smap_filtered, smap_subset = process_smap(dual_smap_frame)

    out_table = pd.concat([smap_filtered, unique_case_aneu_subset, unique_case_cnvs_subset],ignore_index=True)
    out_table['ID_case'] = case_id
    out_table['ID_control'] = control_id
    out_table['Cell type'] = celltype
    out_table = process_subset_calls(out_table)

    all_calls = process_all_calls(smap_subset, aneu_comp_table_indexed, cnv_comp_table_indexed, control_smap_frame)
    all_calls['ID_case'] = case_id
    all_calls['ID_control'] = control_id
    all_calls['Cell type'] = celltype
    cnv_calls, sv_calls = reorder_sheet(all_calls)
    sv_calls_filtered = filter_duplicate_calls(sv_calls)
    cnv_statistics, include_cnv_calls = process_case_and_control_json(control_json, case_json, control_id, case_id)


    with pd.ExcelWriter(out_file) as writer:
        file_handle = extract_path_from_handle(out_file)
        print(f"\nWriting SV calls to results excel file: {out_file}")
        sv_calls_filtered.to_excel(writer, sheet_name='SV_calls',index=False,float_format="%.3f", na_rep='NA')
        print(f"Writing SV calls to csv file: {os.path.join(file_handle,'SV_calls.csv')}")
        write_dataframe_to_csv(df=sv_calls_filtered, filename=os.path.join(file_handle,'SV_calls.csv'))
        print(f"Writing SV calls to docx file: {os.path.join(file_handle,'SV_calls.docx')}")
        generate_docx(sv_calls_filtered, os.path.join(file_handle,'SV_calls'))
        if include_cnv_calls == 0:
            print(f"Writing CNV calls to results excel file: {out_file}")
            cnv_calls.to_excel(writer, sheet_name='CNV_and_Aneuploidy_calls',index=False,float_format="%.3f", na_rep='NA')
            print(f"Writing CNV calls to csv file: {os.path.join(file_handle, 'CNV_and_Aneuploidy_calls.csv')}")
            write_dataframe_to_csv(df=cnv_calls, filename=os.path.join(file_handle, 'CNV_and_Aneuploidy_calls.csv'))
            print(f"Writing CNV calls to docx file: {os.path.join(file_handle, 'CNV_and_Aneuploidy_calls.docx')}")
            generate_docx(cnv_calls, os.path.join(file_handle,'CNV_and_Aneuploidy_calls'))
        print(f"Writing CNV statistics to results excel file: {out_file}")
        cnv_statistics.to_excel(writer, sheet_name='CNV_metrics',float_format="%.3f", na_rep='NA')
        print(f"Writing CNV statistics to csv file: {os.path.join(file_handle,'CNV_metrics.csv')}")
        cnv_statistics.reset_index(inplace=True)
        write_dataframe_to_csv(df=cnv_statistics, filename=os.path.join(file_handle,'CNV_metrics.csv'))
        print(f"Writing CNV statistics to docx file: {os.path.join(file_handle,'CNV_metrics.docx')}")
        generate_docx(cnv_statistics, os.path.join(file_handle,'CNV_metrics'))
        center_excel_cells(writer)

def main():
    parser = argparse.ArgumentParser(
        """This script compares Aneuploidy and CNV calls for samples with corresponding dual variant annotation results. This script will generate a table that contains case-specific SVs, CNVs and aneuploidy calls.""")
    parser.add_argument('--dual_aneuploidy', type=str, help="relative path to dual annotation *_Aneuploidy.txt")
    parser.add_argument('--dual_smap', type=str, help="relative path to dual annotation *_Annotated_SV.smap")
    parser.add_argument('--dual_cnv', type=str, help="relative path to dual annotation *_CNV.txt")
    parser.add_argument('--control_aneuploidy', type=str, help="relative path to control *_Aneuploidy.txt")
    parser.add_argument('--control_cnv', type=str, help="relative path to control *_CNV.txt")
    parser.add_argument('--control_smap', type=str, help="relative path to control *_Annotated_SV.smap")
    parser.add_argument('--control_json', type=str, help="relative path to control report.json containing CNV metrics")
    parser.add_argument('--case_json', type=str, help="relative path to control report.json containing CNV metrics")
    parser.add_argument('--out_file', type=str, help="output file handle")
    parser.add_argument('--case_id', type=str, help="Case ID")
    parser.add_argument('--control_id', type=str, help="Control ID")
    parser.add_argument('--celltype', type=str, help="cell type")
    parser.add_argument('--cnv_overlap_percentage', type=float, nargs='?', const=1, default=0.3, help="maximum reciprocal overlap ratio allowed for CNV calls to be considered unique")
    parser.add_argument('--aneuploidy_overlap_percentage', type=float, nargs='?', const=1, default=0.5,  help="maximum fractional difference allowed for Aneuploidy calls to be considered unique")
    parser.add_argument('--cnv_window', type=int, nargs='?', const=1, default=1000, help="bp window buffer at start and end of CNV calls")
    parser.add_argument('--cnv_stitch_window', type=int, nargs='?', const=1, default=550000, help="bp window to merge neighboring CNV calls")

    # parse command line arguments
    args = parser.parse_args()
    dual_aneuploidy = args.dual_aneuploidy
    dual_smap = args.dual_smap
    dual_cnv = args.dual_cnv
    control_aneuploidy = args.control_aneuploidy
    control_cnv = args.control_cnv
    cnv_window = args.cnv_window
    aneuploidy_overlap_percentage = args.aneuploidy_overlap_percentage
    cnv_overlap_percentage = args.cnv_overlap_percentage
    out_file = args.out_file
    case_id = args.case_id
    control_id = args.control_id
    celltype = args.celltype
    cnv_stitch_window = args.cnv_stitch_window
    control_smap = args.control_smap
    control_json = args.control_json
    case_json = args.case_json
    file_handle = extract_path_from_handle(out_file)
    log_file_handle = os.path.join(file_handle,"{case}_vs_{control}.log".format(case=case_id,control=control_id))
    with open(control_json) as json_file:
        data = json.load(json_file)
    access_dict = data['job']['value']['access']
    solve_dict = data['job']['value']['rescale']
    if os.path.exists(log_file_handle):
        os.remove(log_file_handle)
        print("Previous log file deleted.")
    with PrintLogger(log_file_handle):
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        print(f"Processed Dual VAP results : {timestamp}\n")
        print(f"{access_dict['description']} : {access_dict['value']}")
        print(f"{solve_dict['description']} : {solve_dict['value']}\n")
        compare_calls(dual_aneuploidy, dual_smap, dual_cnv, control_aneuploidy, control_cnv, out_file, case_id, control_id,celltype, control_smap, control_json, case_json, cnv_overlap_percentage, aneuploidy_overlap_percentage, cnv_window, cnv_stitch_window)


if __name__ == "__main__":
    main()